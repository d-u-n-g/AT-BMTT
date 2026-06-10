import tkinter as tk
from tkinter import ttk
from tkinter.scrolledtext import ScrolledText
import random
from tkinter import messagebox, filedialog
import hashlib
import PyPDF2
from docx import Document
from reportlab.pdfgen import canvas

root = tk.Tk()
root.title("Elgamal Nhóm 3")
root.geometry("1200x700")
root.resizable(False, False)

frame_key = ttk.LabelFrame(root, text="🔑 SINH KHÓA")
frame_key.place(x=10, y=10, width=380, height=660)

ttk.Label(frame_key, text="Chọn độ dài bit:").place(x=10, y=10)
combo_bit = ttk.Combobox(
    frame_key,
    values=["128", "256", "512", "1024", "2048"],
    state="readonly"
)
combo_bit.place(x=10, y=35, width=330)
combo_bit.current(0)

ttk.Label(frame_key, text="q:").place(x=10, y=75)
entry_q = ttk.Entry(frame_key)
entry_q.place(x=10, y=95, width=330)

ttk.Label(frame_key, text="a:").place(x=10, y=135)
entry_g = ttk.Entry(frame_key)
entry_g.place(x=10, y=155, width=330)

ttk.Label(frame_key, text="Xa:").place(x=10, y=195)
entry_xa = ttk.Entry(frame_key)
entry_xa.place(x=10, y=215, width=330)

ttk.Label(frame_key, text="Ya:").place(x=10, y=255)
entry_ya = ttk.Entry(frame_key)
entry_ya.place(x=10, y=275, width=330)

btn_auto = ttk.Button(frame_key, text="Sinh tự động")
btn_auto.place(x=10, y=330, width=330, height=35)

btn_reset = ttk.Button(frame_key, text="Làm mới")
btn_reset.place(x=10, y=375, width=330, height=35)

btn_confirm = ttk.Button(frame_key, text="Xác nhận")
btn_confirm.place(x=10, y=420, width=330, height=35)

btn_save_key = ttk.Button(frame_key, text="💾 Lưu khóa")
btn_save_key.place(x=10, y=465, width=330, height=35)

btn_open_key = ttk.Button(frame_key, text="📂 Mở khóa")
btn_open_key.place(x=10, y=510, width=330, height=35)


frame_encrypt = ttk.LabelFrame(root, text="🔒 MÃ HÓA")
frame_encrypt.place(x=400, y=10, width=380, height=660)

ttk.Label(frame_encrypt, text="Bản rõ:").place(x=10, y=10)

txt_plain = ScrolledText(frame_encrypt)
txt_plain.place(x=10, y=35, width=340, height=120)

btn_open_plain = ttk.Button(frame_encrypt, text="📂 Mở file")
btn_open_plain.place(x=10, y=165, width=340)

btn_save_plain_encrypt = ttk.Button(
    frame_encrypt,
    text="💾 Lưu bản rõ"
)
btn_save_plain_encrypt.place(
    x=10,
    y=200,
    width=340
)

ttk.Label(frame_encrypt, text="Xa:").place(x=10, y=245)
entry_enc_xa = ttk.Entry(frame_encrypt)
entry_enc_xa.place(x=10, y=265, width=340)

ttk.Label(frame_encrypt, text="Ya:").place(x=10, y=305)
entry_enc_ya = ttk.Entry(frame_encrypt)
entry_enc_ya.place(x=10, y=325, width=340)

btn_encrypt = ttk.Button(frame_encrypt, text="🔐 Mã hóa")
btn_encrypt.place(x=10, y=375, width=340, height=35)

ttk.Label(frame_encrypt, text="Bản mã:").place(x=10, y=430)

txt_cipher = ScrolledText(frame_encrypt)
txt_cipher.place(x=10, y=455, width=340, height=120)

btn_save_cipher = ttk.Button(frame_encrypt, text="💾 Lưu")
btn_save_cipher.place(x=10, y=590, width=340)


frame_decrypt = ttk.LabelFrame(root, text="🔓 GIẢI MÃ")
frame_decrypt.place(x=790, y=10, width=380, height=660)

ttk.Label(frame_decrypt, text="Bản mã:").place(x=10, y=10)

txt_cipher_input = ScrolledText(frame_decrypt)
txt_cipher_input.place(x=10, y=35, width=340, height=120)

btn_open_cipher = ttk.Button(frame_decrypt, text="📂 Mở file")
btn_open_cipher.place(x=10, y=165, width=340)

ttk.Label(frame_decrypt, text="q:").place(x=10, y=210)
entry_dec_q = ttk.Entry(frame_decrypt)
entry_dec_q.place(x=10, y=230, width=340)

ttk.Label(frame_decrypt, text="Xa:").place(x=10, y=270)
entry_dec_xa = ttk.Entry(frame_decrypt)
entry_dec_xa.place(x=10, y=290, width=340)

btn_decrypt = ttk.Button(frame_decrypt, text="🔓 Giải mã")
btn_decrypt.place(x=10, y=340, width=340, height=35)

ttk.Label(frame_decrypt, text="Bản rõ:").place(x=10, y=400)

txt_plain_output = ScrolledText(frame_decrypt)
txt_plain_output.place(x=10, y=425, width=340, height=120)

btn_save_plain = ttk.Button(frame_decrypt, text="💾 Lưu")
btn_save_plain.place(x=10, y=565, width=340)

def is_prime(n, k=5):
    if n < 2:
        return False
    if n in (2, 3):
        return True
    if n % 2 == 0:
        return False

    r = 0
    d = n - 1
    while d % 2 == 0:
        r += 1
        d //= 2

    for _ in range(k):
        a = random.randint(2, n - 2)
        x = pow(a, d, n)
        if x == 1 or x == n - 1:
            continue
        for _ in range(r - 1):
            x = pow(x, 2, n)
            if x == n - 1:
                break
        else:
            return False

    return True

def generate_prime(bits):

    while True:

        q = random.getrandbits(bits)

        q |= (1 << (bits - 1))
        q |= 1

        if is_prime(q):
            return q

def find_generator(q):
    p = (q - 1) // 2
    while True:
        g = random.randint(2, q - 2)
        if pow(g, 2, q) == 1:
            continue
        if pow(g, p, q) == 1:
            continue
        return g

def auto_generate_key():
    bits = int(combo_bit.get())
    q = generate_prime(bits)
    a = find_generator(q)
    XA = random.randint(2, q - 2)
    YA = pow(a, XA, q)

    entry_q.delete(0, tk.END)
    entry_q.insert(0, str(q))

    entry_g.delete(0, tk.END)
    entry_g.insert(0, str(a))

    entry_xa.delete(0, tk.END)
    entry_xa.insert(0, str(XA))

    entry_ya.delete(0, tk.END)
    entry_ya.insert(0, str(YA))

    entry_enc_xa.delete(0, tk.END)
    entry_enc_xa.insert(0, str(XA))

    entry_enc_ya.delete(0, tk.END)
    entry_enc_ya.insert(0, str(YA))

    print("Đang gán q =", q)
    entry_dec_q.delete(0, tk.END)
    entry_dec_q.insert(0, str(q))

    entry_dec_xa.delete(0, tk.END)
    entry_dec_xa.insert(0, str(XA))

    with open("key_info.txt", "w", encoding="utf-8") as f:
        f.write(f"q={q}\n")
        f.write(f"a={a}\n")
        f.write(f"XA={XA}\n")
        f.write(f"YA={YA}\n")

    messagebox.showinfo("Thông báo", "Sinh khóa thành công!")

def manual_generate_key():
    try:
        q = int(entry_q.get())
        bits = int(combo_bit.get())

        if not is_prime(q):
            messagebox.showerror("Lỗi", "q không phải số nguyên tố")
            return

        if q.bit_length() < bits:

            answer = messagebox.askyesno(
               "Thông báo",
               f"q hiện tại chỉ có {q.bit_length()} bit.\n"
               f"Bạn có muốn hệ thống tự bổ sung lên {bits} bit không?"
            )

            if not answer:
             return

            q |= (1 << (bits - 1))
            q |= 1
            while not is_prime(q):
             q += 2
                

        a = find_generator(q)
        XA = random.randint(2, q - 2)
        YA = pow(a, XA, q)

        entry_q.delete(0, tk.END)
        entry_q.insert(0, str(q))

        entry_g.delete(0, tk.END)
        entry_g.insert(0, str(a))

        entry_xa.delete(0, tk.END)
        entry_xa.insert(0, str(XA))

        entry_ya.delete(0, tk.END)
        entry_ya.insert(0, str(YA))

        entry_enc_xa.delete(0, tk.END)
        entry_enc_xa.insert(0, str(XA))

        entry_enc_ya.delete(0, tk.END)
        entry_enc_ya.insert(0, str(YA))

        entry_dec_q.delete(0, tk.END)
        entry_dec_q.insert(0, str(q))

        entry_dec_xa.delete(0, tk.END)
        entry_dec_xa.insert(0, str(XA))

        with open("key_info.txt", "w", encoding="utf-8") as f:
            f.write(f"q={q}\n")
            f.write(f"a={a}\n")
            f.write(f"XA={XA}\n")
            f.write(f"YA={YA}\n")

        messagebox.showinfo("Thông báo", "Khóa đã được bổ sung!")

    except Exception:
        messagebox.showerror("Lỗi", "Vui lòng nhập q hợp lệ")

def reset_key():
    entry_q.delete(0, tk.END)
    entry_g.delete(0, tk.END)
    entry_xa.delete(0, tk.END)
    entry_ya.delete(0, tk.END)

def save_key_file():
    q = entry_q.get()
    a = entry_g.get()
    XA = entry_xa.get()
    YA = entry_ya.get()

    if q == "" or a == "" or XA == "" or YA == "":
        messagebox.showwarning("Thông báo", "Chưa có khóa để lưu!")
        return

    file_path = filedialog.asksaveasfilename(
        title="Lưu khóa ElGamal",
        defaultextension=".txt",
        filetypes=[
            ("Text File", "*.txt"),
            ("Word File", "*.docx"),
            ("PDF File", "*.pdf")
        ]
    )

    if file_path == "":
        return

    hash_public = hashlib.sha256(
        f"{q}|{a}|{YA}".encode("utf-8")
    ).hexdigest()

    hash_private = hashlib.sha256(
        XA.encode("utf-8")
    ).hexdigest()

    content = (
        "========== KHÓA ELGAMAL ==========\n\n"
        f"q={q}\n"
        f"a={a}\n"
        f"XA={XA}\n"
        f"YA={YA}\n\n"
        f"HASH_PUBLIC={hash_public}\n"
        f"HASH_PRIVATE={hash_private}\n"
    )

    fields = [
        ("q", q),
        ("a", a),
        ("XA", XA),
        ("YA", YA),
        ("HASH_PUBLIC", hash_public),
        ("HASH_PRIVATE", hash_private),
    ]

    if file_path.endswith(".txt"):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write("========== KHÓA ELGAMAL ==========\n\n")
            for key, val in fields:
                f.write(f"{key}={val}\n")

    elif file_path.endswith(".docx"):
        doc = Document()
        doc.add_paragraph("========== KHÓA ELGAMAL ==========")
        for key, val in fields:

            chunk_size = 100
            chunks = [val[i:i+chunk_size] for i in range(0, len(val), chunk_size)]
            if len(chunks) == 1:
                doc.add_paragraph(f"{key}={val}")
            else:
                for idx, chunk in enumerate(chunks):
                    doc.add_paragraph(f"{key}_PART_{idx}={chunk}")
        doc.save(file_path)

    elif file_path.endswith(".pdf"):
        from reportlab.lib.pagesizes import A4
        pdf = canvas.Canvas(file_path, pagesize=A4)
        pdf.setFont("Courier", 8)
        margin = 50
        line_h = 12
        chunk_size = 85  
        y = 800
        pdf.setFont("Helvetica-Bold", 11)
        pdf.drawString(margin, y, "KHOA ELGAMAL")
        y -= 20
        pdf.setFont("Courier", 8)
        for key, val in fields:
            label = f"{key}="
            pdf.drawString(margin, y, label)
            y -= line_h
            for i in range(0, len(val), chunk_size):
                pdf.drawString(margin + 10, y, val[i:i+chunk_size])
                y -= line_h
                if y < 50:
                    pdf.showPage()
                    pdf.setFont("Courier", 8)
                    y = 800
            y -= 4  
        pdf.save()

    messagebox.showinfo("Thông báo", "Lưu khóa thành công!")


def open_key_file():
    file_path = filedialog.askopenfilename(
        title="Mở khóa",
        filetypes=[
            ("Text File", "*.txt"),
            ("Word File", "*.docx"),
            ("PDF File", "*.pdf")
        ]
    )

    if file_path == "":
        return

    try:
        if file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                lines = f.readlines()

        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            lines = [p.text for p in doc.paragraphs]

        elif file_path.endswith(".pdf"):
            raw = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    raw += page.extract_text() or ""
            merged = []
            for raw_line in raw.splitlines():
                raw_line = raw_line.strip()
                if not raw_line:
                    continue
                if "=" in raw_line or raw_line.startswith("=====") or raw_line.startswith("KHOA"):
                    merged.append(raw_line)
                else:
                    if merged:
                        merged[-1] += raw_line
                    else:
                        merged.append(raw_line)
            lines = merged

        else:
            messagebox.showerror("Lỗi", "Định dạng không được hỗ trợ")
            return

        q = a = XA = YA = ""
        old_hash_public = old_hash_private = ""
        parts = {}

        for line in lines:
            line = line.strip()
            if not line or line.startswith("====="):
                continue
            if "=" not in line:
                continue
            key, val = line.split("=", 1)
            key = key.strip()
            val = val.strip()
            if "_PART_" in key:
                base_key = key[:key.index("_PART_")]
                parts.setdefault(base_key, []).append((key, val))
            else:
                if key == "q":
                    q = val
                elif key == "a":
                    a = val
                elif key == "XA":
                    XA = val
                elif key == "YA":
                    YA = val
                elif key == "HASH_PUBLIC":
                    old_hash_public = val
                elif key == "HASH_PRIVATE":
                    old_hash_private = val

        def join_parts(base_key):
            if base_key not in parts:
                return None
            sorted_chunks = sorted(parts[base_key], key=lambda x: int(x[0].split("_PART_")[1]))
            return "".join(v for _, v in sorted_chunks)

        if join_parts("q"):    q  = join_parts("q")
        if join_parts("a"):    a  = join_parts("a")
        if join_parts("XA"):   XA = join_parts("XA")
        if join_parts("YA"):   YA = join_parts("YA")
        if join_parts("HASH_PUBLIC"):  old_hash_public  = join_parts("HASH_PUBLIC")
        if join_parts("HASH_PRIVATE"): old_hash_private = join_parts("HASH_PRIVATE")

        new_hash_public = hashlib.sha256(
            f"{q}|{a}|{YA}".encode("utf-8")
        ).hexdigest()

        new_hash_private = hashlib.sha256(
            XA.encode("utf-8")
        ).hexdigest()

        entry_q.delete(0, tk.END)
        entry_q.insert(0, q)

        entry_g.delete(0, tk.END)
        entry_g.insert(0, a)

        entry_xa.delete(0, tk.END)
        entry_xa.insert(0, XA)

        entry_ya.delete(0, tk.END)
        entry_ya.insert(0, YA)

        entry_enc_xa.delete(0, tk.END)
        entry_enc_xa.insert(0, XA)

        entry_enc_ya.delete(0, tk.END)
        entry_enc_ya.insert(0, YA)

        entry_dec_q.delete(0, tk.END)
        entry_dec_q.insert(0, q)

        entry_dec_xa.delete(0, tk.END)
        entry_dec_xa.insert(0, XA)

        if new_hash_public == old_hash_public and new_hash_private == old_hash_private:
            messagebox.showinfo("Kiểm tra khóa", "✓ Khóa nguyên vẹn")
        elif new_hash_public != old_hash_public and new_hash_private == old_hash_private:
            messagebox.showerror("Kiểm tra khóa", "✗ Khóa công khai đã bị sửa")
        elif new_hash_public == old_hash_public and new_hash_private != old_hash_private:
            messagebox.showerror("Kiểm tra khóa", "✗ Khóa bí mật đã bị sửa")
        else:
            messagebox.showerror("Kiểm tra khóa", "✗ Cả khóa công khai và khóa bí mật đều bị sửa")

    except Exception as e:
        messagebox.showerror("Lỗi", str(e))

def read_file(file_path):
    lower = file_path.lower()

    if lower.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif lower.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    elif lower.endswith(".pdf"):
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text

    else:
        raise Exception("Không hỗ trợ định dạng file")

def open_plain_file():
    file_path = filedialog.askopenfilename(
        title="Mở bản rõ",
        filetypes=[
            ("Text File", "*.txt"),
            ("Word File", "*.docx"),
            ("PDF File", "*.pdf")
        ]
    )

    if not file_path:
        return

    try:
        data = read_file(file_path)
        txt_plain.delete("1.0", tk.END)
        txt_plain.insert(tk.END, data)
    except Exception as e:
        messagebox.showerror("Lỗi", str(e))

def save_plain_encrypt():

    text = txt_plain.get("1.0", tk.END).strip()

    if text == "":
        messagebox.showwarning(
            "Thông báo",
            "Không có nội dung để lưu!"
        )
        return

    file_path = filedialog.asksaveasfilename(
        title="Lưu bản rõ",
        defaultextension=".txt",
        filetypes=[
            ("Text File", "*.txt"),
            ("Word File", "*.docx"),
            ("PDF File", "*.pdf")
        ]
    )

    if file_path == "":
        return

    try:

        if file_path.endswith(".txt"):

            with open(
                file_path,
                "w",
                encoding="utf-8"
            ) as f:

                f.write(text)

        elif file_path.endswith(".docx"):

            doc = Document()
            doc.add_paragraph(text)
            doc.save(file_path)

        elif file_path.endswith(".pdf"):

            pdf = canvas.Canvas(file_path)

            y = 800

            for line in text.split("\n"):

                pdf.drawString(
                    50,
                    y,
                    line
                )

                y -= 20

            pdf.save()

        messagebox.showinfo(
            "Thông báo",
            "Lưu bản rõ thành công!"
        )

    except Exception as e:

        messagebox.showerror(
            "Lỗi",
            str(e)
        )

def encrypt_text():
    try:
        q = int(entry_q.get())    
        a = int(entry_g.get())     
        ya = int(entry_enc_ya.get()) 

        plaintext = txt_plain.get("1.0", tk.END).strip()

        if plaintext == "":
            messagebox.showwarning("Thông báo", "Chưa có bản rõ!")
            return

        cipher_list = []

        for ch in plaintext:
            m = ord(ch)
            k = random.randint(2, q - 2)
            C1 = pow(a, k, q)
            C2 = (m * pow(ya, k, q)) % q
            cipher_list.append(f"{C1},{C2}")

        cipher_text = "|".join(cipher_list)

        txt_cipher.delete("1.0", tk.END)
        txt_cipher.insert(tk.END, cipher_text)

        messagebox.showinfo("Thông báo", "Mã hóa thành công!")

    except Exception as e:
        messagebox.showerror("Lỗi", str(e))

def save_cipher_file():
    cipher = txt_cipher.get("1.0", tk.END).strip()

    if cipher == "":
        messagebox.showwarning("Thông báo", "Chưa có bản mã!")
        return

    file_path = filedialog.asksaveasfilename(
        title="Lưu bản mã",
        defaultextension=".txt",
        filetypes=[
            ("Text File", "*.txt"),
            ("Word File", "*.docx"),
            ("PDF File", "*.pdf")
        ]
    )

    if file_path == "":
        return

    hash_cipher = hashlib.sha256(cipher.encode("utf-8")).hexdigest()

    if file_path.endswith(".txt"):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"CIPHER={cipher}\n")
            f.write(f"HASH_CIPHER={hash_cipher}\n")

    elif file_path.endswith(".docx"):
        doc = Document()
        chunk_size = 200
        doc.add_paragraph("CIPHER_START")
        for i in range(0, len(cipher), chunk_size):
            doc.add_paragraph(cipher[i:i+chunk_size])
        doc.add_paragraph("CIPHER_END")
        doc.add_paragraph(f"HASH_CIPHER={hash_cipher}")
        doc.save(file_path)

    elif file_path.endswith(".pdf"):
        pdf = canvas.Canvas(file_path)
        pdf.setFont("Helvetica", 9)
        margin = 50
        y = 800
        pdf.drawString(margin, y, f"HASH_CIPHER={hash_cipher}")
        y -= 20
        pdf.drawString(margin, y, "CIPHER_START")
        y -= 14
        chunk_size = 90
        for i in range(0, len(cipher), chunk_size):
            pdf.drawString(margin, y, cipher[i:i+chunk_size])
            y -= 13
            if y < 50:
                pdf.showPage()
                pdf.setFont("Helvetica", 9)
                y = 800
        pdf.drawString(margin, y, "CIPHER_END")
        pdf.save()

    messagebox.showinfo("Thông báo", "Lưu bản mã thành công!")

def open_cipher_file():
    file_path = filedialog.askopenfilename(
        title="Mở bản mã",
        filetypes=[
            ("Text File", "*.txt"),
            ("Word File", "*.docx"),
            ("PDF File", "*.pdf")
        ]
    )

    if file_path == "":
        return

    try:
        content = read_file(file_path)
        lines = content.splitlines()

        cipher = ""
        old_hash = ""

        in_cipher_block = False
        cipher_chunks = []

        for line in lines:
            line_stripped = line.strip()
            if line_stripped == "CIPHER_START":
                in_cipher_block = True
            elif line_stripped == "CIPHER_END":
                in_cipher_block = False
                cipher = "".join(cipher_chunks)
            elif in_cipher_block:
                cipher_chunks.append(line_stripped)
            elif line_stripped.startswith("CIPHER="):
                cipher = line_stripped.replace("CIPHER=", "", 1)
            elif line_stripped.startswith("HASH_CIPHER="):
                old_hash = line_stripped.replace("HASH_CIPHER=", "", 1)

        new_hash = hashlib.sha256(cipher.encode("utf-8")).hexdigest()

        txt_cipher_input.delete("1.0", tk.END)
        txt_cipher_input.insert(tk.END, cipher)

        if new_hash == old_hash:
            messagebox.showinfo("Kiểm tra bản mã", "✓ Bản mã nguyên vẹn")
        else:
            messagebox.showerror("Kiểm tra bản mã", "✗ Bản mã đã bị sửa")

    except Exception as e:
        messagebox.showerror("Lỗi", str(e))

def decrypt_text():
    try:
        q = int(entry_dec_q.get())          # q từ khung sinh khóa
        XA = int(entry_dec_xa.get())     # Xa từ khung giải mã

        cipher_text = txt_cipher_input.get("1.0", tk.END).strip()

        if cipher_text == "":
            messagebox.showwarning("Thông báo", "Chưa có bản mã!")
            return

        plaintext = ""
        blocks = cipher_text.split("|")

        for block in blocks:
            C1, C2 = map(int, block.split(","))
            K = pow(C1, XA, q)
            K_inv = pow(K, q - 2, q)
            m = (C2 * K_inv) % q
            plaintext += chr(m)

        txt_plain_output.delete("1.0", tk.END)
        txt_plain_output.insert(tk.END, plaintext)

        messagebox.showinfo("Thông báo", "Giải mã thành công!")

    except Exception as e:
        messagebox.showerror("Lỗi", str(e))

def save_plain_file():
    plaintext = txt_plain_output.get("1.0", tk.END).strip()

    if plaintext == "":
        messagebox.showwarning("Thông báo", "Chưa có bản rõ!")
        return

    file_path = filedialog.asksaveasfilename(
        title="Lưu bản rõ",
        defaultextension=".txt",
        filetypes=[
            ("Text File", "*.txt"),
            ("Word File", "*.docx"),
            ("PDF File", "*.pdf")
        ]
    )

    if file_path == "":
        return

    if file_path.endswith(".txt"):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(plaintext)

    elif file_path.endswith(".docx"):
        doc = Document()
        doc.add_paragraph(plaintext)
        doc.save(file_path)

    elif file_path.endswith(".pdf"):
        pdf = canvas.Canvas(file_path)
        pdf.setFont("Helvetica", 11)
        margin = 50
        max_chars = 90
        y = 800
        for raw_line in plaintext.split("\n"):
            while len(raw_line) > 0:
                chunk = raw_line[:max_chars]
                raw_line = raw_line[max_chars:]
                pdf.drawString(margin, y, chunk)
                y -= 16
                if y < 50:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 11)
                    y = 800
        pdf.save()

    messagebox.showinfo("Thông báo", "Lưu bản rõ thành công!")

btn_auto.config(command=auto_generate_key)
btn_confirm.config(command=manual_generate_key)
btn_reset.config(command=reset_key)
btn_save_key.config(command=save_key_file)
btn_open_key.config(command=open_key_file)
btn_open_plain.config(command=open_plain_file)
btn_save_plain_encrypt.config(command=save_plain_encrypt)
btn_encrypt.config(command=encrypt_text)
btn_save_cipher.config(command=save_cipher_file)
btn_decrypt.config(command=decrypt_text)
btn_open_cipher.config(command=open_cipher_file)
btn_save_plain.config(command=save_plain_file)

root.mainloop()